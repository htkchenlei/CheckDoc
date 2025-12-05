# app.py
from flask import Flask, render_template, request, jsonify
import os
import re
import tempfile
import json
import uuid
from werkzeug.utils import secure_filename as werkzeug_secure_filename
from docx import Document

# --- 配置 ---
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}  # 为简化，暂时只支持 docx。doc 支持需要额外库且复杂。
MAX_CONTENT_LENGTH = 100 * 1024 * 1024  # 增加到 100MB max file size
REGIONS_FILE = 'china_regions.json'

app = Flask(__name__)
app.secret_key = 'your_strong_secret_key_here'  # 更换为强密钥用于生产环境
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# 确保上传文件夹存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def load_regions():
    """从 JSON 文件加载地域名称"""
    if os.path.exists(REGIONS_FILE):
        try:
            with open(REGIONS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # 如果是列表格式，直接返回
                if isinstance(data, list):
                    return data
                # 如果是字典格式，提取所有地域名称
                elif isinstance(data, dict):
                    regions = []
                    # 遍历所有省份
                    for province, cities in data.items():
                        regions.append(province)
                        # 如果城市是字典形式（包含区县），提取所有城市和区县
                        if isinstance(cities, dict):
                            for city, districts in cities.items():
                                regions.append(city)
                                if isinstance(districts, list):
                                    regions.extend(districts)
                        # 如果城市是列表形式（直接列出城市），添加所有城市
                        elif isinstance(cities, list):
                            regions.extend(cities)
                    return regions
                else:
                    app.logger.error(f"JSON file {REGIONS_FILE} format is invalid.")
                    return []
        except (json.JSONDecodeError, IOError) as e:
            app.logger.error(f"Error reading {REGIONS_FILE}: {e}")
            return []
    return []


def save_regions(regions_data):
    """将地域名称保存到 JSON 文件"""
    try:
        # 如果传入的是列表，则转换为字典格式
        if isinstance(regions_data, list):
            # 尝试从现有文件加载结构化数据
            structured_data = {}
            if os.path.exists(REGIONS_FILE):
                try:
                    with open(REGIONS_FILE, 'r', encoding='utf-8') as f:
                        existing_data = json.load(f)
                        if isinstance(existing_data, dict):
                            structured_data = existing_data
                except:
                    pass
            
            # 如果没有现有结构化数据，则使用默认结构
            if not structured_data:
                structured_data = {
                    "省级": [],
                    "市级": [],
                    "区级": []
                }
                
                # 简单分类，实际应用中可能需要更复杂的逻辑
                for region in regions_data:
                    if region not in structured_data["省级"] and region not in structured_data["市级"] and region not in structured_data["区级"]:
                        # 这里只是一个简单的示例，实际应用中可能需要更好的分类方法
                        structured_data["区级"].append(region)
            
            regions_data = structured_data
        # 如果传入的是字典格式，直接使用
        elif isinstance(regions_data, dict):
            # 确保所有必要的键都存在
            required_keys = ["省级", "市级", "区级"]
            for key in required_keys:
                if key not in regions_data:
                    regions_data[key] = []
        
        # 在保存前清理可能存在的无效条目
        for level in ["省级", "市级", "区级"]:
            if level in regions_data:
                # 过滤掉空字符串和仅包含问号的条目
                regions_data[level] = [region for region in regions_data[level] 
                                      if region and not region.isspace() and region != "???" and region != "？？？"]
        
        # 确保使用UTF-8编码并处理可能的编码问题
        with open(REGIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump(regions_data, f, ensure_ascii=False, indent=4)
        return True
    except IOError as e:
        app.logger.error(f"Error saving to {REGIONS_FILE}: {e}")
        return False
    except Exception as e:
        app.logger.error(f"Unexpected error saving to {REGIONS_FILE}: {e}")
        return False


class DocxTextExtractor:
    """专门用于提取 docx 文本和页码信息的类"""

    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.page_breaks = self._find_page_breaks()
        self.full_text = self._extract_full_text()

    def _find_page_breaks(self):
        """估算段落在文档中的页码位置"""
        page_breaks = []
        current_page = 1
        current_line_count = 0
        lines_per_page_estimate = 40  # 这是一个估算值，可根据需要调整

        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text
            if text.strip():  # 忽略空段落
                # 估算段落行数
                lines_in_para = len(text) // 80 + text.count('\n') + 1  # 简单估算
                current_line_count += lines_in_para

                # 检查段落中是否有分页符
                if 'w:br' in paragraph._p.xml and 'type="page"' in paragraph._p.xml:
                    # 找到显式分页符
                    page_breaks.append((i, current_page))
                    current_page += 1
                    current_line_count = lines_in_para  # 重置行数计数

                # 检查是否因行数估算而翻页
                if current_line_count > lines_per_page_estimate:
                    page_breaks.append((i, current_page))
                    current_page += 1
                    current_line_count = lines_in_para  # 重置为当前段落的行数

        return page_breaks

    def _extract_full_text(self):
        """提取完整文本"""
        full_text = []
        for para in self.doc.paragraphs:
            # 保留换行符以便于上下文查找
            full_text.append(para.text + "\n")
        return ''.join(full_text)

    def get_page_number(self, paragraph_index):
        """根据段落索引估算页码"""
        page_num = 1
        for break_index, break_page in self.page_breaks:
            if paragraph_index >= break_index:
                page_num = break_page + 1  # 下一页开始
            else:
                break
        return page_num

    def find_keyword_occurrences(self, keywords):
        """查找关键词并返回其页码和上下文"""
        occurrences = []
        context_length = 50  # 上下文字符数

        # 为了提高查找效率，可以将列表转换为集合进行成员检查
        keywords_set = set(kw.strip() for kw in keywords if kw.strip())

        for kw_index, keyword in enumerate(keywords):
            keyword = keyword.strip()
            if not keyword:
                continue
            start = 0
            while True:
                pos = self.full_text.find(keyword, start)
                if pos == -1:
                    break

                # 找到包含关键词的段落索引
                para_index = 0
                text_pos = 0
                for i, para in enumerate(self.doc.paragraphs):
                    para_end = text_pos + len(para.text) + 1  # +1 for \n
                    if text_pos <= pos < para_end:
                        para_index = i
                        break
                    text_pos = para_end

                # 估算页码
                page_num = self.get_page_number(para_index)

                # 提取上下文
                context_start = max(0, pos - context_length)
                context_end = min(len(self.full_text), pos + len(keyword) + context_length)
                context = self.full_text[context_start:context_end].strip()

                occurrences.append({
                    'keyword': keyword,
                    'page': page_num,
                    'context': context
                })
                start = pos + 1  # 从下一个位置继续查找

        return occurrences


def secure_filename(filename):
    """
    改进的secure_filename函数，支持中文文件名
    """
    # 如果文件名包含中文，我们采用不同的处理方式
    if re.search(r'[\u4e00-\u9fff]', filename):
        # 保留中文字符，但移除危险字符
        # 只保留字母、数字、中文、点号、下划线和连字符
        safe_name = re.sub(r'[^\w\u4e00-\u9fff.-]', '', filename)
        if not safe_name:
            # 如果过滤后为空，使用默认名称
            safe_name = 'upload'

        # 确保文件扩展名被保留
        if '.' in filename:
            parts = filename.rsplit('.', 1)
            if len(parts) == 2:
                extension = parts[1]
                # 清理扩展名中的非法字符
                clean_extension = re.sub(r'[^\w]', '', extension)
                if clean_extension:
                    # 如果安全名称中已有扩展名，先移除它
                    if '.' in safe_name:
                        safe_name = safe_name.rsplit('.', 1)[0]
                    safe_name = f"{safe_name}.{clean_extension}"
                else:
                    # 如果没有有效的扩展名，添加默认的
                    if '.' not in safe_name:
                        safe_name = f"{safe_name}.docx"
    else:
        # 对于非中文文件名，使用原始的secure_filename函数
        safe_name = werkzeug_secure_filename(filename)

    # 特别处理包含引号等特殊字符的文件名
    safe_name = safe_name.replace('"', '').replace("'", "").replace("<", "").replace(">", "")
    return safe_name


def allowed_file(filename):
    """检查文件扩展名是否被允许"""
    try:
        app.logger.info(f"Checking file: {filename}")

        # 检查文件名是否为空
        if not filename:
            app.logger.warning("Empty filename provided")
            return False

        # 检查是否包含点号
        if '.' not in filename:
            app.logger.warning(f"No dot found in filename: {filename}")
            return False

        # 获取文件扩展名
        parts = filename.rsplit('.', 1)
        # print(parts)
        app.logger.info(f"Filename parts: {parts}")

        if len(parts) != 2:
            app.logger.warning(f"Unexpected number of parts in filename: {filename}")
            return False

        extension = parts[1].lower()
        app.logger.info(f"Extension: {extension}, Allowed: {ALLOWED_EXTENSIONS}")

        result = extension in ALLOWED_EXTENSIONS
        app.logger.info(f"File allowed: {result}")
        return result
    except Exception as e:
        # 记录异常但不中断程序
        app.logger.error(f"Error checking file extension for filename '{filename}': {str(e)}")
        return False


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        # 检查是否有文件部分
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        file = request.files['file']

        # 如果用户没有选择文件
        if file.filename == '':
            return jsonify({'success': False, 'message': '没有选择文件'}), 400

        # 检查文件类型和保存
        if file and allowed_file(file.filename):
            # 使用改进的secure_filename函数
            filename = secure_filename(file.filename)
            app.logger.info(f"Original filename: {file.filename}, Secure filename: {filename}")

            # 使用临时文件处理，避免磁盘写入
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1])
            file.save(temp_file.name)
            temp_file_path = temp_file.name
            temp_file.close()  # 关闭文件以便后续读取

            try:
                # 获取关键词类型
                check_type = request.form.get('checkType', 'custom')

                keywords = []
                if check_type == 'custom':
                    keywords_str = request.form.get('keywords', '')
                    if not keywords_str.strip():
                        return jsonify({'success': False, 'message': '请输入至少一个关键词'}), 400
                    keywords = keywords_str.split('，')  # 使用中文逗号分割
                elif check_type == 'china_regions':
                    keywords = load_regions()  # 从 JSON 文件加载地域名
                    # 确保 keywords 是一个列表
                    if keywords is None:
                        keywords = []
                    if not keywords:
                        return jsonify({'success': False, 'message': '地域名称列表为空或加载失败'}), 500
                else:
                    return jsonify({'success': False, 'message': '无效的检查类型'}), 400

                # 提取文本和查找关键词
                if filename.lower().endswith('.docx'):
                    extractor = DocxTextExtractor(temp_file_path)
                    occurrences = extractor.find_keyword_occurrences(keywords)
                    os.unlink(temp_file_path)  # 处理完立即删除临时文件
                    return jsonify(
                        {'success': True, 'filename': filename, 'occurrences': occurrences, 'checkType': check_type})

                else:
                    os.unlink(temp_file_path)
                    return jsonify({'success': False, 'message': '不支持的文件类型。请上传 .docx 文件。'}), 400

            except Exception as e:
                # 确保即使出错也删除临时文件
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                app.logger.error(f"Error processing file {filename}: {e}")
                return jsonify({'success': False, 'message': f'处理文件时出错: {str(e)}'}), 500

        else:
            return jsonify({'success': False, 'message': '不支持的文件类型。请上传 .docx 文件。'}), 400
            
    except Exception as e:
        app.logger.error(f"Error in upload_file: {e}")
        # 确保返回JSON格式的错误信息，避免前端解析错误
        return jsonify({'success': False, 'message': f'上传文件时发生错误: {str(e)}'}), 500


# --- 地域名称管理 API ---

@app.route('/api/regions', methods=['GET'])
def get_regions():
    """获取所有地域名称"""
    regions = load_regions()
    return jsonify({'success': True, 'regions': regions})


@app.route('/api/regions', methods=['POST'])
def add_region():
    """添加一个新的地域名称"""
    data = request.get_json()
    new_region = data.get('region', '').strip()

    if not new_region:
        return jsonify({'success': False, 'message': '地域名称不能为空'}), 400

    regions = load_regions()
    # 确保 regions 是一个列表
    if regions is None:
        regions = []
    if new_region in regions:
        return jsonify({'success': False, 'message': '该地域名称已存在'}), 400

    regions.append(new_region)
    if save_regions(regions):
        return jsonify({'success': True, 'message': '地域名称添加成功', 'region': new_region})
    else:
        return jsonify({'success': False, 'message': '保存地域名称失败'}), 500


@app.route('/api/regions/<region>', methods=['PUT'])
def update_region(region):
    """修改一个地域名称"""
    data = request.get_json()
    updated_region = data.get('updatedRegion', '').strip()

    if not updated_region:
        return jsonify({'success': False, 'message': '新地域名称不能为空'}), 400

    regions = load_regions()
    # 确保 regions 是一个列表
    if regions is None:
        regions = []
    if region not in regions:
        return jsonify({'success': False, 'message': '要修改的地域名称不存在'}), 404

    if updated_region in regions:
        return jsonify({'success': False, 'message': '新地域名称已存在'}), 400

    index = regions.index(region)
    regions[index] = updated_region
    if save_regions(regions):
        return jsonify(
            {'success': True, 'message': '地域名称修改成功', 'oldRegion': region, 'newRegion': updated_region})
    else:
        return jsonify({'success': False, 'message': '保存地域名称失败'}), 500


@app.route('/api/regions/<region>', methods=['DELETE'])
def delete_region(region):
    """删除一个地域名称"""
    regions = load_regions()
    # 确保 regions 是一个列表
    if regions is None:
        regions = []
    if region not in regions:
        return jsonify({'success': False, 'message': '要删除的地域名称不存在'}), 404

    regions.remove(region)
    if save_regions(regions):
        return jsonify({'success': True, 'message': '地域名称删除成功', 'region': region})
    else:
        return jsonify({'success': False, 'message': '保存地域名称失败'}), 500


@app.route('/api/regions/structured', methods=['GET'])
def get_structured_regions():
    """获取结构化的地域名称"""
    regions = load_regions_structured()
    return jsonify({'success': True, 'regions': regions})


@app.route('/api/regions/<level>', methods=['POST'])
def add_region_by_level(level):
    """根据级别添加地域名称"""
    app.logger.info(f"Adding region to level: {level}")
    
    regions = load_regions_structured()
    app.logger.info(f"Current regions: {regions}")

    if level not in regions:
        app.logger.error(f"Invalid level: {level}")
        return jsonify({'success': False, 'message': '无效的级别'}), 400

    try:
        # 尝试多种方式获取数据
        data = request.get_json(force=True)
        if data and 'name' in data:
            new_region = data['name']
        else:
            # 如果JSON解析失败或没有name字段，尝试从表单数据中获取
            new_region = request.form.get('name', '')
        
        app.logger.info(f"Raw received new region name: '{new_region}' (type: {type(new_region)})")
        
        # 确保new_region是字符串类型
        new_region = str(new_region)
    except Exception as e:
        app.logger.error(f"Error getting region name: {e}")
        new_region = ''

    # 去除首尾空白字符
    new_region = new_region.strip()
    app.logger.info(f"Processed region name: '{new_region}'")

    if not new_region:
        app.logger.error("Empty region name provided")
        return jsonify({'success': False, 'message': '地域名称不能为空'}), 400

    if new_region in regions[level]:
        app.logger.error(f"Region '{new_region}' already exists in level '{level}'")
        return jsonify({'success': False, 'message': '该地域名称已存在'}), 400

    regions[level].append(new_region)
    app.logger.info(f"Added region '{new_region}' to level '{level}'. New list: {regions[level]}")

    if save_regions(regions):
        app.logger.info(f"Successfully saved regions. Returning region name: '{new_region}'")
        return jsonify({'success': True, 'message': '地域名称添加成功', 'name': new_region})
    else:
        app.logger.error("Failed to save regions")
        return jsonify({'success': False, 'message': '保存地域名称失败'}), 500


@app.route('/api/regions/<level>/<name>', methods=['DELETE'])
def delete_region_by_level(level, name):
    """根据级别删除地域名称"""
    try:
        # 尝试获取请求中的JSON数据
        data = request.get_json()
        if data is None:
            # 如果没有JSON数据，尝试从表单数据中获取密码
            password = request.form.get('password', '')
        else:
            password = data.get('password', '')
    except:
        password = ''

    # 验证密码
    if password != '123456':
        return jsonify({'success': False, 'message': '密码错误'}), 403

    try:
        regions = load_regions_structured()

        if level not in regions:
            return jsonify({'success': False, 'message': '无效的级别'}), 400

        if name not in regions[level]:
            return jsonify({'success': False, 'message': '地域名称不存在'}), 404

        regions[level].remove(name)

        if save_regions(regions):
            return jsonify({'success': True, 'message': '地域名称删除成功', 'name': name})
        else:
            return jsonify({'success': False, 'message': '保存地域名称失败'}), 500
    except Exception as e:
        app.logger.error(f"Error deleting region: {e}")
        return jsonify({'success': False, 'message': f'删除过程中发生错误: {str(e)}'}), 500


def load_regions_structured():
    """从 JSON 文件加载结构化的地域名称"""
    if os.path.exists(REGIONS_FILE):
        try:
            with open(REGIONS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # 确保返回的是字典格式
                if isinstance(data, dict):
                    return data
                else:
                    app.logger.error(f"JSON file {REGIONS_FILE} format is invalid (not a dict).")
                    return {"省级": [], "市级": [], "区级": []}
        except (json.JSONDecodeError, IOError) as e:
            app.logger.error(f"Error reading {REGIONS_FILE}: {e}")
            return {"省级": [], "市级": [], "区级": []}
    else:
        # 如果文件不存在，创建一个带有默认数据的文件
        default_regions = {
            "省级": ["北京", "天津", "上海", "重庆", "香港", "澳门", "内蒙古", "广西", "西藏", "宁夏", "新疆",
                     "河北", "山西", "辽宁", "吉林", "黑龙江", "江苏", "浙江", "安徽", "福建", "江西", "山东",
                     "河南", "湖北", "湖南", "广东", "海南", "四川", "贵州", "云南", "陕西", "甘肃", "青海"],
            "市级": ["石家庄", "唐山", "张家口", "太原", "沈阳", "大连", "辽阳", "长春", "松原", "延边",
                     "哈尔滨", "齐齐哈尔", "南京", "无锡", "徐州", "常州", "苏州", "南通", "连云港", "淮安",
                     "盐城", "扬州", "镇江", "泰州", "宿迁", "杭州", "宁波", "温州", "嘉兴", "湖州", "绍兴",
                     "金华", "衢州", "舟山", "台州", "丽水", "合肥", "六安", "亳州", "福州", "厦门", "南昌",
                     "赣州", "济南", "青岛", "淄博", "郑州", "开封", "洛阳", "焦作", "武汉", "黄石", "十堰",
                     "宜昌", "襄阳", "鄂州", "荆门", "孝感", "荆州", "黄冈", "咸宁", "随州", "恩施", "仙桃",
                     "潜江", "天门", "神农架林区", "长沙", "株洲", "湘潭", "岳阳", "广州", "深圳", "肇庆",
                     "惠州", "梅州", "海口", "三亚", "成都", "甘孜", "凉山", "贵阳", "黔西南", "黔东南", "黔南",
                     "昆明", "西安", "榆林", "兰州", "西宁"],
            "区级": ["东城", "西城", "朝阳", "丰台", "石景山", "海淀", "门头沟", "房山", "通州", "顺义",
                     "昌平", "大兴", "怀柔", "平谷", "密云", "延庆", "和平", "河东", "河西", "南开", "河北",
                     "红桥", "东丽", "西青", "津南", "北辰", "武清", "宝坻", "滨海新", "宁河", "静海", "蓟州",
                     "黄浦", "徐汇", "长宁", "静安", "普陀", "虹口", "杨浦", "闵行", "宝山", "嘉定", "浦东",
                     "金山", "松江", "青浦", "奉贤", "崇明", "万州", "涪陵", "渝中", "大渡口", "江北", "沙坪坝",
                     "九龙坡", "南岸", "北碚", "綦江", "大足", "渝北", "巴南", "黔江", "长寿", "江津", "合川",
                     "永川", "南川", "璧山", "铜梁", "潼南", "荣昌", "开州", "梁平", "武隆", "城口", "丰都",
                     "垫江", "忠县", "云阳", "奉节", "巫山", "巫溪"]
        }
        save_regions(default_regions)
        return default_regions

if __name__ == '__main__':
    app.run(debug=True, port=5003)#生产环境中请设置 debug=False